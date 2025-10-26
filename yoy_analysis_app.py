# sales_dashboard_all.py
import streamlit as st
import pandas as pd
import numpy as np
# import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import io
import os

st.set_page_config(page_title="PepsiCo Campaign Analysis Dashboard - 14 May 2025 to 14 August 2025", layout="wide")

# ---------------- Helpers ----------------
def human_format(num):
    if pd.isna(num):
        return ""
    try:
        num = float(num)
    except:
        return str(num)
    if abs(num) >= 1_000_000:
        return f'{num/1_000_000:.1f}M'
    elif abs(num) >= 1_000:
        return f'{num/1_000:.0f}K'
    else:
        return str(int(num))

def human_currency(num):
    if pd.isna(num):
        return ""
    try:
        num = float(num)
    except:
        return str(num)
    if abs(num) >= 1_000_000:
        return f'R{num/1_000_000:.1f}M'
    elif abs(num) >= 1_000:
        return f'R{num/1_000:.0f}K'
    else:
        return f'R{num:.0f}'

def add_bolded_paragraph(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    run.font.name = 'Aptos'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
    run.font.size = Pt(9)
    run2 = p.add_run(normal_text)
    run2.font.name = 'Aptos'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
    run2.font.size = Pt(9)
    return p

def save_fig_to_bytes(fig, fmt="png", dpi=300):
    buf = io.BytesIO()
    fig.savefig(buf, format=fmt, dpi=dpi, bbox_inches='tight')
    buf.seek(0)
    return buf

def doc_style_setup(doc):
    try:
        style = doc.styles['Normal']
        style.font.name = 'Aptos'
        style.font.size = Pt(9)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
    except Exception:
        pass

# ---------------- YOY Analysis ----------------
def yoy_analysis_page():
    st.header("YOY Analysis")    

    # Read Excel
    try:
        df = pd.read_excel('YOY Analysis.xlsx', sheet_name=0, header=1)
        # Normalize expected columns
        df.columns = [c.strip() for c in df.columns]
        expected = ['Product Description', 'QTY Sold Prior Year', 'QTY Sold CAMPAIGN PERIOD', 'Increase in sales from Prior Year AVE']
        if set(expected).issubset(set(df.columns)):
            df = df[expected]
        else:
            # Try to map first 4 columns if names differ
            df = df.iloc[:, :4]
            df.columns = expected
        df = df[df['Product Description'].apply(lambda x: isinstance(x, str))]
    except FileNotFoundError:
        st.error("File `YOY Analysis.xlsx` not found in current directory.")
        return
    except Exception as e:
        st.error(f"Error reading YOY file: {e}")
        return

    # Metrics
    increase_count = (df['Increase in sales from Prior Year AVE'] > 0).sum()
    decrease_count = (df['Increase in sales from Prior Year AVE'] < 0).sum()
    avg_increase = df['Increase in sales from Prior Year AVE'].mean()
    top_grower = df.loc[df['Increase in sales from Prior Year AVE'].idxmax()]
    top_decliner = df.loc[df['Increase in sales from Prior Year AVE'].idxmin()]

    total_prior = df['QTY Sold Prior Year'].sum()
    total_campaign = df['QTY Sold CAMPAIGN PERIOD'].sum()
    total_growth = (total_campaign - total_prior) / total_prior if total_prior != 0 else np.nan

    # Chart 1: YOY pct change
    products = df['Product Description']
    increases = df['Increase in sales from Prior Year AVE']
    colors = ['green' if x > 0 else 'red' for x in increases]
    yoy_pct = increases * 100

    fig1, ax1 = plt.subplots(figsize=(12, 6), constrained_layout=True)
    bars = ax1.bar(products, yoy_pct, color=colors)
    ax1.set_ylabel('Increase in Sales from Prior Year (%)')
    ax1.set_title('YOY Sales Change by Product')
    ax1.tick_params(axis='x', rotation=90)
    label_margin = max(2, abs(yoy_pct.max()) * 0.04)
    ymax = yoy_pct.max() + label_margin
    ymin = min(0, yoy_pct.min())
    ax1.set_ylim(ymin, ymax + label_margin * 1.5)
    for bar, value in zip(bars, yoy_pct):
        label = f"{value:.1f}%"
        label_y = bar.get_height() + label_margin * 0.7
        ax1.text(bar.get_x() + bar.get_width()/2, label_y, label, ha='center', va='bottom', fontsize=8, color='black', clip_on=False)
    st.pyplot(fig1)

    # Chart 2: Prior vs Campaign volumes
    x = np.arange(len(products))
    width = 0.35
    prior = df['QTY Sold Prior Year']
    campaign = df['QTY Sold CAMPAIGN PERIOD']

    fig2, ax2 = plt.subplots(figsize=(12, 6), constrained_layout=True)
    bars1 = ax2.bar(x - width/2, prior, width, label='Prior Year', color='lightgray')
    bars2 = ax2.bar(x + width/2, campaign, width, label='Campaign Period', color='royalblue')
    ax2.set_ylabel('Quantity Sold')
    ax2.set_title('YOY Sales Comparison by Product')
    ax2.set_xticks(x)
    ax2.set_xticklabels(products, rotation=90)
    ax2.legend()
    all_heights = np.concatenate([prior.values, campaign.values])
    label_margin2 = max(2000, all_heights.max() * 0.04)
    ymax2 = all_heights.max() + label_margin2
    ax2.set_ylim(0, ymax2 + label_margin2 * 1.5)
    for i, bar in enumerate(bars1):
        value = prior.iloc[i]
        label_y = bar.get_height() + label_margin2 * 0.7
        ax2.text(bar.get_x() + bar.get_width()/2, label_y, human_format(value), ha='center', va='bottom', fontsize=8, color='black', clip_on=False)
    for i, bar in enumerate(bars2):
        value = campaign.iloc[i]
        label_y = bar.get_height() + label_margin2 * 0.7
        ax2.text(bar.get_x() + bar.get_width()/2, label_y, human_format(value), ha='center', va='bottom', fontsize=8, color='black', clip_on=False)
    st.pyplot(fig2)

    # Key insights display
    st.subheader("Key Findings")
    st.markdown(f"- {increase_count} products increased, {decrease_count} decreased; average change {avg_increase:.2%}.")
    st.markdown(f"- Top grower: **{top_grower['Product Description']}** ({top_grower['Increase in sales from Prior Year AVE']:.2%}).")
    st.markdown(f"- Top decliner: **{top_decliner['Product Description']}** ({top_decliner['Increase in sales from Prior Year AVE']:.2%}).")
    st.markdown(f"- Total prior: {human_format(total_prior)} → Total campaign: {human_format(total_campaign)} ({total_growth:.2%}).")

    # Generate DOCX
    if st.button("Generate YOY Executive Summary (.docx)"):
        try:
            # Save figs to bytes
            img1 = save_fig_to_bytes(fig1)
            img2 = save_fig_to_bytes(fig2)

            doc = Document()
            doc_style_setup(doc)
            title = doc.add_heading('Executive Summary', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle = doc.add_paragraph('Year-on-Year Sales Analysis')
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in subtitle.runs:
                run.font.size = Pt(14)
                run.font.name = 'Aptos'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            intro = (
                "This executive summary provides a concise overview of key findings from the year-on-year (YOY) sales analysis. "
                "The analysis compares product sales between the prior year and the recent campaign period, highlighting significant trends and changes in performance metrics."
            )
            p_intro = doc.add_paragraph(intro)
            for run in p_intro.runs:
                run.font.name = 'Aptos'
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            # Key Findings
            doc.add_heading('Key Findings', level=1)
            key_findings = [
                f"• Out of {len(df)} products analyzed, {increase_count} showed an increase in sales, while {decrease_count} experienced a decline during the campaign period compared to the prior year.",
                f"• The average change in sales across all products was {avg_increase:.2%}.",
                f"• The product with the highest sales growth was {top_grower['Product Description']} with a {top_grower['Increase in sales from Prior Year AVE']:.2%} increase.",
                f"• The product with the largest decline was {top_decliner['Product Description']} with a {top_decliner['Increase in sales from Prior Year AVE']:.2%} decrease.",
                f"• Overall, total sales increased from {human_format(total_prior)} units in the prior year to {human_format(total_campaign)} units in the campaign period, representing a total growth of {total_growth:.2%}."
            ]
            for finding in key_findings:
                p = doc.add_paragraph(finding)
                for run in p.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            # Charts
            doc.add_heading('YOY Sales Change by Product', level=1)
            doc.add_picture(img1, width=Inches(6))
            explanation_chart1 = [
                ("Observations: ", f"This chart displays YOY % change. {increase_count} products grew, {decrease_count} declined. Top grower: {top_grower['Product Description']}."),
                ("Interpretation: ", f"Green bars indicate positive momentum; red indicate decline. Avg change: {avg_increase:.1%}."),
                ("Opportunities: ", "Sustain successful tactics and investigate declines.")
            ]
            for bold, normal in explanation_chart1:
                add_bolded_paragraph(doc, bold, " " + normal)
            doc.add_paragraph()

            doc.add_heading('YOY Sales Comparison by Product', level=1)
            doc.add_picture(img2, width=Inches(6))
            explanation_chart2 = [
                ("Observations: ", f"Prior vs campaign volumes. Top campaign product: {df.iloc[df['QTY Sold CAMPAIGN PERIOD'].idxmax()]['Product Description']}."),
                ("Interpretation: ", "Higher campaign volumes indicate campaign impact."),
                ("Opportunities: ", "Replicate successful campaign tactics.")
            ]
            for bold, normal in explanation_chart2:
                add_bolded_paragraph(doc, bold, " " + normal)
            doc.add_paragraph()

            doc.add_heading('Summary', level=1)
            p_summary = doc.add_paragraph(
                "Overall positive trend observed. Focus on sustaining momentum for high performers and diagnose declines for targeted interventions."
            )
            for run in p_summary.runs:
                run.font.name = 'Aptos'
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

            # Output docx to BytesIO and provide download
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            st.success("YOY executive summary ready.")
            st.download_button(
                label="Download YOY Executive Summary (.docx)",
                data=doc_io.getvalue(),
                file_name="YOY_Executive_Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Failed to generate YOY document: {e}")

# ---------------- Prior Periods ----------------
def prior_periods_page():
    st.header("Campaign Prior Periods Analysis")
    

    # Read Excel
    try:
        df = pd.read_excel('Prior Periods.xlsx', header=1)
        df.columns = [c.strip() for c in df.columns]
        # locate expected columns (fuzzy)
        col_prior = next((c for c in df.columns if 'Prior Year' in c), None)
        col_feb_may = next((c for c in df.columns if 'Feb' in c or '14 Feb' in c), None)
        col_campaign = next((c for c in df.columns if 'CAMPAIGN' in c.upper()), None)
        col_increase = next((c for c in df.columns if 'Increase' in c), None)

        if not (col_prior and col_feb_may and col_campaign and col_increase and 'Product Description' in df.columns):
            st.warning("Could not detect all expected columns automatically. Please ensure the Excel file structure matches the expected layout.")
        df = df[df['Product Description'].apply(lambda x: isinstance(x, str))]
        # Compute avg prior months if possible
        if col_prior and col_feb_may:
            df['Avg Prior Months'] = df[[col_prior, col_feb_may]].mean(axis=1)
        else:
            df['Avg Prior Months'] = np.nan
    except FileNotFoundError:
        st.error("File `Prior Periods.xlsx` not found in current directory.")
        return
    except Exception as e:
        st.error(f"Error reading Prior Periods file: {e}")
        return

    # Chart 1
    x = np.arange(len(df['Product Description']))
    width = 0.35
    campaign_vals = df[col_campaign] / 10000 if col_campaign else np.zeros(len(df))
    avg_prior_vals = df['Avg Prior Months'] / 10000

    fig1, ax1 = plt.subplots(figsize=(12,6), constrained_layout=True)
    bars1 = ax1.bar(x - width/2, campaign_vals, width, label='Campaign Period Sales')
    bars2 = ax1.bar(x + width/2, avg_prior_vals, width, label='Avg of Prior Months')
    ax1.set_ylabel('Quantity Sold (10,000s)')
    ax1.set_xticks(x)
    ax1.set_xticklabels(df['Product Description'], rotation=90)
    ax1.legend()
    
    # Fix: Convert to numpy arrays before using flatten
    all_heights = np.concatenate([campaign_vals.values, avg_prior_vals.values])
    max_height = all_heights.max() if len(all_heights)>0 else 0
    ymax = max_height * 1.12 if max_height>0 else 1
    ax1.set_ylim(0, ymax)
    
    for i, bar in enumerate(bars1):
        value = df[col_campaign].iloc[i] if col_campaign else np.nan
        label_y = bar.get_height() + ymax * 0.01
        ax1.text(bar.get_x() + bar.get_width()/2, label_y, human_format(value), ha='center', va='bottom', fontsize=8, color='black', clip_on=True)
    for i, bar in enumerate(bars2):
        value = df['Avg Prior Months'].iloc[i]
        label_y = bar.get_height() + ymax * 0.01
        ax1.text(bar.get_x() + bar.get_width()/2, label_y, human_format(value), ha='center', va='bottom', fontsize=8, color='black', clip_on=True)
    st.pyplot(fig1)

    # Key findings for Chart 1
    st.subheader("Key Findings - Campaign vs Prior Average")
    products_above_avg = df[df[col_campaign] > df['Avg Prior Months']]['Product Description'].tolist() if col_campaign else []
    products_below_avg = df[df[col_campaign] < df['Avg Prior Months']]['Product Description'].tolist() if col_campaign else []
    st.markdown(f"- {len(products_above_avg)} out of {len(df)} products achieved higher sales during the campaign than their prior average.")
    st.markdown(f"- Products above average likely benefited from campaign activities.")
    st.markdown(f"- Focus on replicating successful tactics and investigating underperformers.")

    # Chart 2: increase %
    increase_pct = df[col_increase] * 100 if col_increase else np.zeros(len(df))
    colors = ['green' if v >= 0 else 'red' for v in increase_pct]

    fig2, ax2 = plt.subplots(figsize=(12,6), constrained_layout=True)
    bars = ax2.bar(df['Product Description'], increase_pct, color=colors)
    ax2.set_ylabel('Sales Increase (%)')
    ax2.set_title('Sales Increase During Campaign vs. Avg of Prior Months')
    ax2.tick_params(axis='x', rotation=90)
    ymin = min(0, np.nanmin(increase_pct)) if len(increase_pct)>0 else 0
    ymax = max(0, np.nanmax(increase_pct)) if len(increase_pct)>0 else 0
    y_range = ymax - ymin if (ymax - ymin) != 0 else 1
    ax2.set_ylim(ymin - y_range * 0.12, ymax + y_range * 0.12)
    for bar, value in zip(bars, increase_pct):
        label = f'{value:.1f}%'
        if value >= 0:
            label_y = value + y_range * 0.01
            va = 'bottom'
        else:
            label_y = value - y_range * 0.01
            va = 'top'
        ax2.text(bar.get_x() + bar.get_width()/2, label_y, label, ha='center', va=va, fontsize=8, clip_on=True)
    st.pyplot(fig2)

    # Key findings
    increase_count = (df[col_increase] > 0).sum() if col_increase else 0
    decrease_count = (df[col_increase] < 0).sum() if col_increase else 0
    avg_increase = df[col_increase].mean() if col_increase else np.nan
    top_grower = df.loc[df[col_increase].idxmax()] if col_increase else None
    top_decliner = df.loc[df[col_increase].idxmin()] if col_increase else None
    top_campaign = df.iloc[df[col_campaign].idxmax()] if col_campaign else None
    top_avg_prior = df.iloc[df['Avg Prior Months'].idxmax()] if 'Avg Prior Months' in df.columns else None

    st.subheader("Key Findings - Sales Increase Analysis")
    st.markdown(f"- {increase_count} products increased vs avg prior months, {decrease_count} decreased.")
    if top_grower is not None and col_increase:
        st.markdown(f"- Top grower vs prior average: **{top_grower['Product Description']}** (+{top_grower[col_increase]:.1%}).")
    st.markdown(f"- {increase_count} products experienced growth, while {decrease_count} declined.")
    st.markdown(f"- Green bars indicate positive campaign impact.")
    st.markdown(f"- Focus on doubling down on growth products.")

    # Generate DOCX
    if st.button("Generate Prior Periods Summary (.docx)"):
        try:
            img1 = save_fig_to_bytes(fig1)
            img2 = save_fig_to_bytes(fig2)

            doc = Document()
            doc_style_setup(doc)
            title = doc.add_heading('Executive Summary: Prior Periods Analysis', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            intro = (
                "This executive summary provides a concise overview of key findings from the campaign sales analysis. "
                "The analysis compares product sales during the campaign period to the average of prior months."
            )
            p_intro = doc.add_paragraph(intro)
            for run in p_intro.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Key Findings', level=1)
            key_findings = [
                f"- Out of {len(df)} products analyzed, {increase_count} showed an increase in sales during the campaign period compared to the average of prior months, while {decrease_count} experienced a decline.",
                f"- The average change in sales across all products was {avg_increase:.2%}." if not pd.isna(avg_increase) else ""
            ]
            for finding in key_findings:
                if finding:
                    p = doc.add_paragraph(finding)
                    for run in p.runs:
                        run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Campaign Period Sales vs. Average of Prior Months', level=1)
            doc.add_picture(img1, width=Inches(6))
            expl1 = [
                ("Observations: ",
                 f"{len(products_above_avg)} out of {len(df)} products achieved higher sales during the campaign than their prior average."),
                ("Interpretation: ", "Products above average likely benefited from campaign activities."),
                ("Opportunities: ", "Replicate successful tactics and investigate underperformers.")
            ]
            for bold, normal in expl1:
                add_bolded_paragraph(doc, bold, " " + normal)
            doc.add_paragraph()

            doc.add_heading('Sales Increase During Campaign vs. Avg of Prior Months', level=1)
            doc.add_picture(img2, width=Inches(6))
            expl2 = [
                ("Observations: ", f"{increase_count} products experienced growth, while {decrease_count} declined."),
                ("Interpretation: ", "Green bars indicate positive campaign impact."),
                ("Opportunities: ", "Double down on growth products.")
            ]
            for bold, normal in expl2:
                add_bolded_paragraph(doc, bold, " " + normal)
            doc.add_paragraph()

            doc.add_heading('Summary', level=1)
            summary_text = (
                "The data indicates a generally positive trend in sales performance during the campaign period. "
                "Stakeholders should focus on sustaining the momentum for top-performing products while investigating the causes behind declining products."
            )
            p_summary = doc.add_paragraph(summary_text)
            for run in p_summary.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

            doc_io = io.BytesIO()
            doc.save(doc_io); doc_io.seek(0)
            st.success("Prior Periods summary generated.")
            st.download_button(
                label="Download Prior Periods Summary (.docx)",
                data=doc_io.getvalue(),
                file_name="Prior_Periods_Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"Failed to generate Prior Periods document: {e}")

# ---------------- Category Analysis ----------------
def category_analysis_page():
    st.header("Campaign Category Share Analysis")

    # Data (hard-coded)
    data = {
        "Product Description": [
            "BOKOMO CORN FLAKES CEREALS ORIGINAL 1KG",
            "BOKOMO TRADITIONAL OATS 1KG",
            "LIQUI FRUIT 2L",
            "SIMBA CHIPS 120G",
            "WEET BIX CEREALS BOX 450G",
            "WEET BIX CEREALS BOX 900G",
            "WELLINGTONS SWEET CHILLI SAUCE 700ML",
            "WELLINGTONS TOMATO SAUCE 700ml",
            "WHITE STAR INSTANT MAIZE PORRIDGE 1KG",
            "WHITE STAR SUPER MAIZE MEAL MAIZE BAG 2.5KG",
            "WHITE STAR M/MEAL 10KG"
        ],
        "PepsiCo Campaign Category Share": [0.80, 0.14, 0.82, 0.74, 0.79, 0.83, 0.62, 0.45, 0.33, 0.09, 0.06],
        "Competitor Category Share": [0.20, 0.86, 0.18, 0.26, 0.21, 0.17, 0.38, 0.55, 0.67, 0.91, 0.94],
        "Pre-Campaign PepsiCo": [0.76, 0.10, 0.85, 0.73, 0.77, 0.80, 0.69, 0.47, 0.30, 0.12, 0.07],
        "% Change": [0.05, 0.04, -0.03, 0.01, 0.02, 0.02, -0.06, -0.02, 0.03, -0.03, -0.01]
    }
    df = pd.DataFrame(data)

    x = np.arange(len(df["Product Description"]))
    width = 0.35

    # Chart 1
    fig1, ax1 = plt.subplots(figsize=(12,6), constrained_layout=True)
    bars1 = ax1.bar(x - width/2, df["PepsiCo Campaign Category Share"]*100, width, label='PepsiCo', color='royalblue')
    bars2 = ax1.bar(x + width/2, df["Competitor Category Share"]*100, width, label='Competitors', color='lightgray')
    ax1.set_ylabel('Category Share (%)')
    ax1.set_title('Category Share During Campaign Period by Product')
    ax1.set_xticks(x); ax1.set_xticklabels(df["Product Description"], rotation=90)
    ax1.set_ylim(0,100)
    ax1.legend()
    for bar in bars1:
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1, f"{bar.get_height():.0f}%", ha='center', va='bottom', fontsize=8)
    for bar in bars2:
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1, f"{bar.get_height():.0f}%", ha='center', va='bottom', fontsize=8)
    st.pyplot(fig1)

    # Chart 2
    fig2, ax2 = plt.subplots(figsize=(12,6), constrained_layout=True)
    bars = ax2.bar(x, df["Pre-Campaign PepsiCo"]*100, width=0.5, label='Pre-Campaign PepsiCo', color='royalblue')
    ax2.set_ylabel('Pre-Campaign PepsiCo Share (%)')
    ax2.set_title('Pre-Campaign PepsiCo Share by Product')
    ax2.set_xticks(x); ax2.set_xticklabels(df["Product Description"], rotation=90)
    ymax = (df["Pre-Campaign PepsiCo"]*100).max() + 15
    ax2.set_ylim(0, ymax)
    for i, bar in enumerate(bars):
        bar_height = bar.get_height()
        pct_change = df["% Change"].iloc[i] * 100
        ax2.text(bar.get_x() + bar.get_width()/2, bar_height/2, f"{pct_change:+.0f}%", ha='center', va='center', fontsize=10, color='white', fontweight='bold')
        ax2.text(bar.get_x() + bar.get_width()/2, bar_height + 1, f"{bar_height:.0f}%", ha='center', va='bottom', fontsize=8, color='black')
    st.pyplot(fig2)

    # Chart 3
    fig3, ax3 = plt.subplots(figsize=(12,6), constrained_layout=True)
    width3 = 0.2
    bars_campaign = ax3.bar(x - width3*1.5, df["PepsiCo Campaign Category Share"]*100, width3, label='PepsiCo (Campaign)', color='royalblue')
    bars_pre = ax3.bar(x - width3/2, df["Pre-Campaign PepsiCo"]*100, width3, label='PepsiCo (Pre-Campaign)', color='deepskyblue')
    bars_comp = ax3.bar(x + width3/2, df["Competitor Category Share"]*100, width3, label='Competitor (Campaign)', color='lightgray')
    ax3.set_ylabel('Category Share (%)'); ax3.set_title('Category Share Comparison by Product (Campaign vs Pre-Campaign)')
    ax3.set_xticks(x); ax3.set_xticklabels(df["Product Description"], rotation=90)
    all_heights = np.concatenate([
        df["PepsiCo Campaign Category Share"]*100,
        df["Pre-Campaign PepsiCo"]*100,
        df["Competitor Category Share"]*100
    ])
    ymax = all_heights.max() + 12
    ax3.set_ylim(0, ymax)
    for bars in [bars_campaign, bars_pre, bars_comp]:
        for bar in bars:
            ax3.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1, f"{bar.get_height():.0f}%", ha='center', va='bottom', fontsize=8)
    ax3.legend()
    st.pyplot(fig3)

    # Key findings
    top_gain = df.iloc[df["% Change"].idxmax()]
    top_loss = df.iloc[df["% Change"].idxmin()]
    avg_pepsico_share = df["PepsiCo Campaign Category Share"].mean()
    avg_competitor_share = df["Competitor Category Share"].mean()
    num_gain = (df["% Change"] > 0).sum()
    num_loss = (df["% Change"] < 0).sum()

    st.subheader("Key Findings")
    st.markdown(f"- PepsiCo avg share {avg_pepsico_share:.0%}, competitors {avg_competitor_share:.0%}.")
    st.markdown(f"- {num_gain} products increased share, {num_loss} decreased. Largest gain: **{top_gain['Product Description']}**.")

    if st.button("Generate Category Share Executive Summary (.docx)"):
        try:
            img1 = save_fig_to_bytes(fig1)
            img2 = save_fig_to_bytes(fig2)
            img3 = save_fig_to_bytes(fig3)

            doc = Document(); doc_style_setup(doc)
            title = doc.add_heading('Executive Summary: Category Share Analysis', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            intro = (
                "This executive summary provides a concise overview of key findings from the category share analysis. "
                "The analysis compares PepsiCo and competitor category shares during the campaign period."
            )
            p_intro = doc.add_paragraph(intro)
            for run in p_intro.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Key Findings', level=1)
            kfs = [
                f"• PepsiCo's average category share during the campaign was {avg_pepsico_share:.0%}, compared to competitors' {avg_competitor_share:.0%}.",
                f"• {num_gain} products increased their PepsiCo category share during the campaign, while {num_loss} saw a decrease.",
                f"• The largest gain in share was for {top_gain['Product Description']} (+{top_gain['% Change']:.0%}), while the largest loss was for {top_loss['Product Description']} ({top_loss['% Change']:.0%})."
            ]
            for finding in kfs:
                p = doc.add_paragraph(finding)
                for run in p.runs:
                    run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Category Share During Campaign Period', level=1)
            doc.add_picture(img1, width=Inches(6))
            add_bolded_paragraph(doc, "Observations: ", f"PepsiCo held an average share of {avg_pepsico_share:.0%}.")
            doc.add_paragraph()

            doc.add_heading('Pre-Campaign PepsiCo Share', level=1)
            doc.add_picture(img2, width=Inches(6))
            add_bolded_paragraph(doc, "Observations: ", "Pre-campaign baseline and % change are shown on the bars.")
            doc.add_paragraph()

            doc.add_heading('Category Share Comparison (Campaign vs Pre-Campaign)', level=1)
            doc.add_picture(img3, width=Inches(6))
            add_bolded_paragraph(doc, "Observations: ", f"{num_gain} products improved PepsiCo share vs pre-campaign.")
            doc.add_paragraph()

            doc.add_heading('Summary', level=1)
            p_summary = doc.add_paragraph(
                "PepsiCo maintained a strong category share in most products during the campaign period. Focus on sustaining growth where observed."
            )
            for run in p_summary.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

            doc_io = io.BytesIO(); doc.save(doc_io); doc_io.seek(0)
            st.success("Category Share summary generated.")
            st.download_button(
                "Download Category Share Executive Summary (.docx)",
                data=doc_io.getvalue(),
                file_name="Category_Share_Executive_Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Failed to generate Category document: {e}")

# ---------------- Campaign Units Analysis ----------------
def campaign_units_page():
    st.header("Campaign Units Analysis (Pre, During, Post)")

    data = {
        "Product Description": [
            "BOKOMO CORN FLAKES CEREALS ORIGINAL 1KG",
            "BOKOMO TRADITIONAL OATS 1KG",
            "LIQUI FRUIT 2L",
            "SIMBA CHIPS 120G",
            "WEET BIX CEREALS BOX 450G",
            "WEET BIX CEREALS BOX 900G",
            "WELLINGTONS SWEET CHILLI SAUCE 700ML",
            "WELLINGTONS TOMATO SAUCE 700ml",
            "WHITE STAR INSTANT MAIZE PORRIDGE 1KG",
            "WHITE STAR SUPER MAIZE MEAL MAIZE BAG 2.5KG",
            "WHITE STAR M/MEAL 10KG"
        ],
        "Pre-Campaign Units (6wks)": [41134, 1116, 43481, 139366, 16784, 21149, 1091, 6336, 34407, 1168, 9259],
        "Campaign Units/Week": [47657, 2146, 53466, 153274, 20570, 21155, 1445, 8003, 36302, 1412, 9409],
        "Post-Campaign Units/Week": [37502, 2115, 44185, 113034, 18148, 17783, 849, 4889, 30976, 1084, 7371],
        "% Change (Campaign vs Pre)": [0.16, 0.92, 0.23, 0.10, 0.23, 0.0003, 0.32, 0.26, 0.06, 0.21, 0.02],
        "% Change (Campaign vs Post)": [-0.21, -0.01, -0.17, -0.26, -0.12, -0.16, -0.41, -0.39, -0.15, -0.23, -0.22]
    }
    df = pd.DataFrame(data)

    x = np.arange(len(df["Product Description"]))
    width = 0.25

    # Chart 1
    fig1, ax1 = plt.subplots(figsize=(12,6), constrained_layout=True)
    bars_pre = ax1.bar(x - width, df["Pre-Campaign Units (6wks)"], width, label='Pre-Campaign', color='lightgray')
    bars_camp = ax1.bar(x, df["Campaign Units/Week"], width, label='Campaign', color='royalblue')
    bars_post = ax1.bar(x + width, df["Post-Campaign Units/Week"], width, label='Post-Campaign', color='deepskyblue')
    ax1.set_ylabel('Units Sold')
    ax1.set_title('Units Sold per Product: Pre-, During-, and Post-Campaign')
    ax1.set_xticks(x); ax1.set_xticklabels(df["Product Description"], rotation=90)
    ax1.legend()
    all_heights = np.concatenate([df["Pre-Campaign Units (6wks)"].values, df["Campaign Units/Week"].values, df["Post-Campaign Units/Week"].values])
    ymax = all_heights.max() * 1.15
    ax1.set_ylim(0, ymax)
    for bars in [bars_pre, bars_camp, bars_post]:
        for bar in bars:
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + ymax*0.01, human_format(bar.get_height()), ha='center', va='bottom', fontsize=8)
    st.pyplot(fig1)

    # Key findings for Chart 1
    st.subheader("Key Findings - Units Sold Comparison")
    st.markdown(f"- This chart compares units sold per product before, during, and after the campaign.")
    st.markdown(f"- Campaign period generally drove higher weekly sales across the portfolio.")
    st.markdown(f"- Focus on strategies to extend the positive effects of campaigns.")

    # Chart 2
    fig2, ax2 = plt.subplots(figsize=(12,6), constrained_layout=True)
    colors = ['green' if v >= 0 else 'red' for v in df["% Change (Campaign vs Pre)"]]
    bars = ax2.bar(df["Product Description"], np.array(df["% Change (Campaign vs Pre)"])*100, color=colors)
    ax2.set_ylabel('% Change (Campaign vs Pre)')
    ax2.set_title('% Change in Units Sold: Campaign vs Pre-Campaign')
    ax2.tick_params(axis='x', rotation=90)
    ymin = min(0, (np.array(df["% Change (Campaign vs Pre)"])*100).min())
    ymax = max(0, (np.array(df["% Change (Campaign vs Pre)"])*100).max())
    y_range = ymax - ymin if (ymax - ymin) != 0 else 1
    ax2.set_ylim(ymin - y_range*0.12, ymax + y_range*0.12)
    for bar, value in zip(bars, np.array(df["% Change (Campaign vs Pre)"])*100):
        label = f'{value:.0f}%'
        label_y = value + y_range*0.01 if value >= 0 else value - y_range*0.01
        va = 'bottom' if value >= 0 else 'top'
        ax2.text(bar.get_x() + bar.get_width()/2, label_y, label, ha='center', va=va, fontsize=8)
    st.pyplot(fig2)

    # Key findings for Chart 2
    st.subheader("Key Findings - Campaign vs Pre-Campaign")
    num_gain_pre = (df["% Change (Campaign vs Pre)"] > 0).sum()
    num_loss_pre = (df["% Change (Campaign vs Pre)"] < 0).sum()
    top_gain_pre = df.iloc[df["% Change (Campaign vs Pre)"].idxmax()]
    st.markdown(f"- {num_gain_pre} products experienced an increase in units sold during the campaign compared to pre-campaign.")
    st.markdown(f"- Products with strong positive change likely benefited from campaign activities.")
    st.markdown(f"- Replicate successful tactics from top-performing products like {top_gain_pre['Product Description']}.")

    # Chart 3
    fig3, ax3 = plt.subplots(figsize=(12,6), constrained_layout=True)
    colors = ['green' if v >= 0 else 'red' for v in df["% Change (Campaign vs Post)"]]
    bars = ax3.bar(df["Product Description"], np.array(df["% Change (Campaign vs Post)"])*100, color=colors)
    ax3.set_ylabel('% Change (Campaign vs Post)')
    ax3.set_title('% Change in Units Sold: Campaign vs Post-Campaign')
    ax3.tick_params(axis='x', rotation=45)
    ymin = min(0, (np.array(df["% Change (Campaign vs Post)"])*100).min())
    ymax = max(0, (np.array(df["% Change (Campaign vs Post)"])*100).max())
    y_range = ymax - ymin if (ymax - ymin) != 0 else 1
    ax3.set_ylim(ymin - y_range*0.12, ymax + y_range*0.12)
    for bar, value in zip(bars, np.array(df["% Change (Campaign vs Post)"])*100):
        label = f'{value:.0f}%'
        label_y = value + y_range*0.01 if value >= 0 else value - y_range*0.01
        va = 'bottom' if value >= 0 else 'top'
        ax3.text(bar.get_x() + bar.get_width()/2, label_y, label, ha='center', va=va, fontsize=8)
    st.pyplot(fig3)

    # Key findings
    num_gain_post = (df["% Change (Campaign vs Post)"] > 0).sum()
    num_loss_post = (df["% Change (Campaign vs Post)"] < 0).sum()
    top_gain_post = df.iloc[df["% Change (Campaign vs Post)"].idxmax()]
    top_loss_post = df.iloc[df["% Change (Campaign vs Post)"].idxmin()]
    avg_change_pre = df["% Change (Campaign vs Pre)"].mean()
    avg_change_post = df["% Change (Campaign vs Post)"].mean()

    st.subheader("Key Findings - Campaign vs Post-Campaign")
    st.markdown(f"- {num_gain_post} products maintained or increased sales post-campaign compared to the campaign period.")
    st.markdown(f"- Most products saw a drop in sales after the campaign.")
    st.markdown(f"- Develop post-campaign plans to sustain gains.")

    # Summary findings
    st.subheader("Summary Findings")
    st.markdown(f"- Campaign period drove higher weekly sales for most products, but these gains were not always sustained post-campaign.")
    st.markdown(f"- Focus on strategies to extend the positive effects of campaigns beyond the campaign period.")

    if st.button("Generate Units Executive Summary (.docx)"):
        try:
            img1 = save_fig_to_bytes(fig1)
            img2 = save_fig_to_bytes(fig2)
            img3 = save_fig_to_bytes(fig3)

            doc = Document(); doc_style_setup(doc)
            title = doc.add_heading('Executive Summary: Pre-, During-, and Post-Campaign Units Analysis', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            intro = (
                "This executive summary provides a concise overview of key findings from the units analysis before, during, and after the campaign."
            )
            p_intro = doc.add_paragraph(intro)
            for run in p_intro.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Key Findings', level=1)
            kfs = [
                f"• Most products saw an increase in weekly units sold during the campaign compared to the pre-campaign period.",
                f"• The largest campaign gain was for {df.iloc[df['% Change (Campaign vs Pre)'].idxmax()]['Product Description']} (+{df['% Change (Campaign vs Pre)'].max():.0%})."
            ]
            for finding in kfs:
                p = doc.add_paragraph(finding)
                for run in p.runs:
                    run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Units Sold per Product: Pre-, During-, and Post-Campaign', level=1)
            doc.add_picture(img1, width=Inches(6))
            for bold, normal in [
                ("Observations: ", "This chart compares units sold per product before, during, and after the campaign."),
                ("Interpretation: ", "Campaign period generally drove higher weekly sales across the portfolio."),
                ("Opportunities: ", "Focus on strategies to extend the positive effects of campaigns.")
            ]:
                add_bolded_paragraph(doc, bold, " " + normal)
            doc.add_paragraph()

            doc.add_heading('% Change in Units Sold: Campaign vs Pre-Campaign', level=1)
            doc.add_picture(img2, width=Inches(6))
            for bold, normal in [
                ("Observations: ", f"{num_gain_pre} products experienced an increase in units sold during the campaign compared to pre-campaign."),
                ("Interpretation: ", "Products with strong positive change likely benefited from campaign activities."),
                ("Opportunities: ", "Replicate successful tactics from top-performing products.")
            ]:
                add_bolded_paragraph(doc, bold, " " + normal)
            doc.add_paragraph()

            doc.add_heading('% Change in Units Sold: Campaign vs Post-Campaign', level=1)
            doc.add_picture(img3, width=Inches(6))
            for bold, normal in [
                ("Observations: ", f"{num_gain_post} products maintained or increased sales post-campaign compared to the campaign period."),
                ("Interpretation: ", "Most products saw a drop in sales after the campaign."),
                ("Opportunities: ", "Develop post-campaign plans to sustain gains.")
            ]:
                add_bolded_paragraph(doc, bold, " " + normal)
            doc.add_paragraph()

            doc.add_heading('Summary', level=1)
            p_summary = doc.add_paragraph(
                "Campaign period drove higher weekly sales for most products, but these gains were not always sustained post-campaign."
            )
            for run in p_summary.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

            doc_io = io.BytesIO(); doc.save(doc_io); doc_io.seek(0)
            st.success("Units executive summary ready.")
            st.download_button(
                "Download Units Executive Summary (.docx)",
                data=doc_io.getvalue(),
                file_name="Pre_During_Post_Campaign_Executive_Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Failed to generate Units document: {e}")

# ---------------- Campaign Sales Amount Analysis ----------------
def campaign_sales_amount_page():
    st.header("Campaign Sales Amount Analysis (Pre, During, Post)")
    data = {
        "Product Description": [
            "BOKOMO CORN FLAKES CEREALS ORIGINAL 1KG",
            "BOKOMO TRADITIONAL OATS 1KG",
            "LIQUI FRUIT 2L",
            "SIMBA CHIPS 120G",
            "WEET BIX CEREALS BOX 450G",
            "WEET BIX CEREALS BOX 900G",
            "WELLINGTONS SWEET CHILLI SAUCE 700ML",
            "WELLINGTONS TOMATO SAUCE 700ml",
            "WHITE STAR INSTANT MAIZE PORRIDGE 1KG",
            "WHITE STAR SUPER MAIZE MEAL MAIZE BAG 2.5KG",
            "WHITE STAR M/MEAL 10KG"
        ],
        "Pre-Campaign Sales (6wks)": [
            2146039.02, 40101.00, 1959100.98, 2458327.73, 486957.01, 1119207.82,
            53829.88, 194247.98, 993367.85, 49663.65, 1215487.70
        ],
        "Campaign Sales/Week": [
            2512772.12, 81147.63, 2403505.63, 2695830.06, 595658.98, 1127400.03,
            72150.90, 258569.12, 1051719.19, 59890.38, 1154819.48
        ],
        "Post-Campaign Sales/Week": [
            1910801.08, 76227.17, 2006715.47, 2047509.76, 511282.47, 959973.53,
            44785.86, 163818.96, 882593.02, 44569.48, 852939.35
        ],
        "% Change (Campaign vs Pre)": [
            0.17, 1.02, 0.23, 0.10, 0.22, 0.01, 0.34, 0.33, 0.06, 0.21, -0.05
        ],
        "% Change (Campaign vs Post)": [
            -0.24, -0.06, -0.17, -0.24, -0.14, -0.15, -0.38, -0.37, -0.16, -0.26, -0.26
        ]
    }
    df = pd.DataFrame(data)

    x = np.arange(len(df["Product Description"]))
    width = 0.35

    # Chart 1
    fig1, ax1 = plt.subplots(figsize=(12,6), constrained_layout=True)
    bars_pre = ax1.bar(x - width, df["Pre-Campaign Sales (6wks)"], width, label='Pre-Campaign', color='lightgray')
    bars_camp = ax1.bar(x, df["Campaign Sales/Week"], width, label='Campaign', color='royalblue')
    bars_post = ax1.bar(x + width, df["Post-Campaign Sales/Week"], width, label='Post-Campaign', color='deepskyblue')
    ax1.set_ylabel('Sales Amount (ZAR)')
    ax1.set_title('Sales Amount per Product: Pre-, During-, and Post-Campaign')
    ax1.set_xticks(x); ax1.set_xticklabels(df["Product Description"], rotation=90)
    ax1.legend()
    all_heights = np.concatenate([df["Pre-Campaign Sales (6wks)"].values, df["Campaign Sales/Week"].values, df["Post-Campaign Sales/Week"].values])
    ymax = all_heights.max() * 1.18
    ax1.set_ylim(0, ymax)
    for bars in [bars_pre, bars_camp, bars_post]:
        for bar in bars:
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + ymax*0.01, human_currency(bar.get_height()), ha='center', va='bottom', fontsize=8, rotation=90)
    st.pyplot(fig1)

    # Key findings for Chart 1
    st.subheader("Key Findings - Sales Amount Comparison")
    st.markdown(f"- This chart compares sales amounts per product before, during, and after the campaign.")
    st.markdown(f"- Campaign period generally drove higher weekly sales amounts across the portfolio.")
    st.markdown(f"- Focus on strategies to extend the positive effects of campaigns.")

    # Chart 2
    fig2, ax2 = plt.subplots(figsize=(12,6), constrained_layout=True)
    colors = ['green' if v >= 0 else 'red' for v in df["% Change (Campaign vs Pre)"]]
    bars = ax2.bar(df["Product Description"], np.array(df["% Change (Campaign vs Pre)"])*100, color=colors)
    ax2.set_ylabel('% Change (Campaign vs Pre Sales)')
    ax2.set_title('% Change in Sales Amount: Campaign vs Pre-Campaign')
    ax2.tick_params(axis='x', rotation=90)
    ymin = min(0, (np.array(df["% Change (Campaign vs Pre)"])*100).min())
    ymax = max(0, (np.array(df["% Change (Campaign vs Pre)"])*100).max())
    y_range = ymax - ymin if (ymax - ymin) != 0 else 1
    ax2.set_ylim(ymin - y_range*0.12, ymax + y_range*0.12)
    for bar, value in zip(bars, np.array(df["% Change (Campaign vs Pre)"])*100):
        label = f'{value:.0f}%'
        label_y = value + (y_range*0.01 if value >= 0 else -y_range*0.01)
        va = 'bottom' if value >= 0 else 'top'
        ax2.text(bar.get_x() + bar.get_width()/2, label_y, label, ha='center', va=va, fontsize=8)
    st.pyplot(fig2)

    # Key findings for Chart 2
    st.subheader("Key Findings - Campaign vs Pre-Campaign Sales")
    num_gain_pre = (df["% Change (Campaign vs Pre)"] > 0).sum()
    top_gain_pre = df.iloc[df["% Change (Campaign vs Pre)"].idxmax()]
    st.markdown(f"- {num_gain_pre} products experienced an increase in sales amount during the campaign compared to pre-campaign.")
    st.markdown(f"- Products with strong positive change likely benefited from campaign activities.")
    st.markdown(f"- Replicate successful tactics from top-performing products like {top_gain_pre['Product Description']}.")

    # Chart 3
    fig3, ax3 = plt.subplots(figsize=(12,6), constrained_layout=True)
    colors = ['green' if v >= 0 else 'red' for v in df["% Change (Campaign vs Post)"]]
    bars = ax3.bar(df["Product Description"], np.array(df["% Change (Campaign vs Post)"])*100, color=colors)
    ax3.set_ylabel('% Change (Campaign vs Post Sales)')
    ax3.set_title('% Change in Sales Amount: Campaign vs Post-Campaign')
    ax3.tick_params(axis='x', rotation=90)
    ymin = min(0, (np.array(df["% Change (Campaign vs Post)"])*100).min())
    ymax = max(0, (np.array(df["% Change (Campaign vs Post)"])*100).max())
    y_range = ymax - ymin if (ymax - ymin) != 0 else 1
    ax3.set_ylim(ymin - y_range*0.12, ymax + y_range*0.12)
    for bar, value in zip(bars, np.array(df["% Change (Campaign vs Post)"])*100):
        label = f'{value:.0f}%'
        label_y = value + (y_range*0.01 if value >= 0 else -y_range*0.01)
        va = 'bottom' if value >= 0 else 'top'
        ax3.text(bar.get_x() + bar.get_width()/2, label_y, label, ha='center', va=va, fontsize=8)
    st.pyplot(fig3)

    # Key findings
    num_gain_post = (df["% Change (Campaign vs Post)"] > 0).sum()
    st.subheader("Key Findings - Campaign vs Post-Campaign Sales")
    st.markdown(f"- {num_gain_post} products maintained or increased sales post-campaign compared to the campaign period.")
    st.markdown(f"- Many products saw declines post-campaign.")
    st.markdown(f"- Develop post-campaign plans to sustain gains.")

    # Summary findings
    st.subheader("Summary Findings")
    st.markdown(f"- Campaign increased sales for many products but not all gains were sustained post-campaign.")
    st.markdown(f"- Focus on strategies to extend the positive effects of campaigns beyond the campaign period.")

    if st.button("Generate Sales Amount Executive Summary (.docx)"):
        try:
            img1 = save_fig_to_bytes(fig1)
            img2 = save_fig_to_bytes(fig2)
            img3 = save_fig_to_bytes(fig3)

            doc = Document(); doc_style_setup(doc)
            title = doc.add_heading('Executive Summary: Pre-, During-, and Post-Campaign Sales Amount Analysis', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            intro = (
                "This executive summary provides a concise overview of key findings from the sales amount analysis before, during, and after the campaign."
            )
            p_intro = doc.add_paragraph(intro)
            for run in p_intro.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Key Findings', level=1)
            kfs = [
                f"• Most products saw an increase in weekly sales amount during the campaign compared to the pre-campaign period.",
                f"• The largest campaign gain was for {df.iloc[df['% Change (Campaign vs Pre)'].idxmax()]['Product Description']} (+{df['% Change (Campaign vs Pre)'].max():.0%})."
            ]
            for finding in kfs:
                p = doc.add_paragraph(finding)
                for run in p.runs:
                    run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

            doc.add_heading('Sales Amount per Product: Pre-, During-, and Post-Campaign', level=1)
            doc.add_picture(img1, width=Inches(6))
            add_bolded_paragraph(doc, "Observations: ", "Most products saw a notable increase in sales during the campaign compared to the pre-campaign period.")
            doc.add_paragraph()

            doc.add_heading('% Change in Sales Amount: Campaign vs Pre-Campaign', level=1)
            doc.add_picture(img2, width=Inches(6))
            add_bolded_paragraph(doc, "Observations: ", f"{num_gain_pre} products experienced an increase in sales amount during the campaign compared to pre-campaign.")
            doc.add_paragraph()

            doc.add_heading('% Change in Sales Amount: Campaign vs Post-Campaign', level=1)
            doc.add_picture(img3, width=Inches(6))
            add_bolded_paragraph(doc, "Observations: ", "Many products saw declines post-campaign.")
            doc.add_paragraph()

            doc.add_heading('Summary', level=1)
            p_summary = doc.add_paragraph("Campaign increased sales for many products but not all gains were sustained post-campaign.")
            for run in p_summary.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

            doc_io = io.BytesIO(); doc.save(doc_io); doc_io.seek(0)
            st.success("Sales amount executive summary ready.")
            st.download_button(
                "Download Sales Amount Executive Summary (.docx)",
                data=doc_io.getvalue(),
                file_name="Pre_During_Post_Campaign_Sales_Executive_Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Failed to generate Sales Amount document: {e}")

# ---------------- Demographics ----------------
def demographics_page():
    st.header("Shopper Demographics")
    # Data input (hard-coded)
    data = {
        "Product": [
            "BOKOMO CORN FLAKES CEREALS ORIGINAL 1KG",
            "BOKOMO TRADITIONAL OATS 1KG",
            "LIQUI FRUIT 2L",
            "SIMBA CHIPS 120G",
            "WEET BIX CEREALS BOX 450G",
            "WEET BIX CEREALS BOX 900G",
            "WELLINGTONS SWEET CHILLI SAUCE 700ML",
            "WELLINGTONS TOMATO SAUCE 700ml",
            "WHITE STAR INSTANT MAIZE PORRIDGE 1KG",
            "WHITE STAR SUPER MAIZE MEAL MAIZE BAG 2.5KG",
            "WHITE STAR M/MEAL 10KG"
        ],
        "Mon": [10,12,8,9,12,11,8,10,12,16,9],
        "Tue": [14,13,12,14,15,15,14,15,15,16,17],
        "Wed": [13,14,12,11,12,12,11,10,13,10,16],
        "Thu": [18,19,15,14,15,16,14,13,17,12,14],
        "Fri": [22,16,22,18,20,20,17,16,21,17,18],
        "Sat": [18,14,18,18,17,17,18,21,18,16,17],
        "Sun": [4,13,12,16,8,9,18,16,5,12,9],
        "Female": [84,79,66,60,76,74,64,68,79,62,57],
        "Male": [16,21,34,40,24,26,36,32,21,38,43],
        "0-18": [4,5,7,6,4,6,8,9,2,6,5],
        "18-24": [6,6,3,5,6,3,3,3,7,6,5],
        "25-34": [30,17,18,23,28,20,16,17,32,22,18],
        "35-44": [32,40,31,30,32,33,33,35,33,26,35],
        "45-54": [17,18,25,23,18,23,24,20,16,25,23],
        "55-64": [6,8,10,9,8,11,10,9,7,9,10],
        "65+": [4,6,6,4,4,5,6,6,2,7,5]
    }
    df = pd.DataFrame(data)

    days = ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]
    day_means = df[days].mean()
    day_means = day_means / day_means.sum() * 100

    genders = ["Female","Male"]
    gender_means = df[genders].mean()
    gender_means = gender_means / gender_means.sum() * 100

    df["0-24"] = df["0-18"] + df["18-24"]
    age_groups = ["0-24","25-34","35-44","45-54","55-64","65+"]
    age_means = df[age_groups].mean()
    age_means = age_means / age_means.sum() * 100

    top_day = day_means.idxmax(); top_day_pct = day_means.max()
    low_day = day_means.idxmin(); low_day_pct = day_means.min()
    female_pct = gender_means['Female']; male_pct = gender_means['Male']
    top_age_group = age_means.idxmax(); top_age_pct = age_means.max()

    # Plots
    def plot_day_of_week():
        fig, ax = plt.subplots(figsize=(8,4.5), constrained_layout=True)
        bars = ax.bar(days, day_means, color=plt.cm.Paired.colors[:len(days)])
        ax.set_ylabel('Share of Shoppers (%)'); ax.set_title('Shoppers by Day of Week (Average % Split)')
        ax.set_ylim(0, max(day_means)*1.15)
        ax.tick_params(axis='x', rotation=0)
        for bar in bars:
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5, f"{bar.get_height():.1f}%", ha='center', va='bottom', fontsize=8)
        return fig

    def plot_gender_pie():
        fig, ax = plt.subplots(figsize=(5,5), constrained_layout=True)
        ax.pie(gender_means, labels=gender_means.index, autopct='%1.0f%%', colors=['#6baed6','#fd8d3c'], startangle=90, textprops={'fontsize': 10})
        ax.set_title('Gender Breakdown (Average % Split)')
        return fig

    def plot_age_groups():
        fig, ax = plt.subplots(figsize=(9,4.5), constrained_layout=True)
        bar_colors = plt.cm.Set2.colors[:len(age_groups)]
        bars = ax.bar(age_groups, age_means, color=bar_colors)
        ax.set_ylabel('Share of Shoppers (%)'); ax.set_title('Age Breakdown (Average % Split)')
        ax.set_ylim(0, max(age_means)*1.15)
        for bar in bars:
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5, f"{bar.get_height():.1f}%", ha='center', va='bottom', fontsize=8)
        return fig

    col1, col2 = st.columns([2,1])
    with col1:
        st.subheader("Shoppers by Day of Week")
        fig_day = plot_day_of_week()
        st.pyplot(fig_day)
        st.markdown(f"**Key Findings - Day of Week:**")
        st.markdown(f"- Shopper activity peaked on **{top_day}** ({top_day_pct:.1f}%) and was lowest on **{low_day}** ({low_day_pct:.1f}%).")
        st.markdown(f"- Friday had the highest activity ({top_day_pct:.1f}%).")
    with col2:
        st.subheader("Gender Breakdown")
        fig_gender = plot_gender_pie()
        st.pyplot(fig_gender)
        st.markdown(f"**Key Findings - Gender:**")
        st.markdown(f"- Female shoppers represented **{female_pct:.1f}%** of the total, with males at **{male_pct:.1f}%**.")
        st.markdown(f"- Female shoppers made up {female_pct:.1f}% of shoppers.")

    st.subheader("Age Breakdown")
    fig_age = plot_age_groups()
    st.pyplot(fig_age)
    st.markdown(f"**Key Findings - Age:**")
    st.markdown(f"- The largest age group was **{top_age_group}** ({top_age_pct:.1f}%).")

    st.markdown("### Summary Findings")
    st.markdown(f"- The demographic analysis reveals patterns that should inform timing, targeting, and messaging of future campaigns.")
    st.markdown(f"- Focus on peak shopping days like **{top_day}** for campaign launches.")
    st.markdown(f"- Tailor messaging to resonate with the dominant **{top_age_group}** age group.")

    if st.button("Generate Demographics Executive Summary (.docx)"):
        try:
            img_day = save_fig_to_bytes(fig_day)
            img_gender = save_fig_to_bytes(fig_gender)
            img_age = save_fig_to_bytes(fig_age)

            doc = Document(); doc_style_setup(doc)
            title = doc.add_heading('Executive Summary: Shopper Demographics During Campaign', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            intro = (
                "This executive summary provides a concise overview of shopper demographics during the campaign period."
            )
            p_intro = doc.add_paragraph(intro)
            for run in p_intro.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Key Findings', level=1)
            key_findings = [
                f"• Shopper activity peaked on {top_day} ({top_day_pct:.1f}%) and was lowest on {low_day} ({low_day_pct:.1f}%).",
                f"• Female shoppers represented {female_pct:.1f}% of the total, with males at {male_pct:.1f}%.",
                f"• The largest age group was {top_age_group} ({top_age_pct:.1f}%)."
            ]
            for finding in key_findings:
                p = doc.add_paragraph(finding)
                for run in p.runs:
                    run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
            doc.add_paragraph()

            doc.add_heading('Shoppers by Day of Week', level=1)
            doc.add_picture(img_day, width=Inches(6))
            add_bolded_paragraph(doc, "Observations: ", f"Friday had the highest activity ({top_day_pct:.1f}%).")
            doc.add_paragraph()

            doc.add_heading('Gender Breakdown', level=1)
            doc.add_picture(img_gender, width=Inches(4))
            add_bolded_paragraph(doc, "Observations: ", f"Female shoppers made up {female_pct:.1f}% of shoppers.")
            doc.add_paragraph()

            doc.add_heading('Age Breakdown', level=1)
            doc.add_picture(img_age, width=Inches(6))
            add_bolded_paragraph(doc, "Observations: ", f"The largest age group was {top_age_group} ({top_age_pct:.1f}%).")
            doc.add_paragraph()

            doc.add_heading('Summary', level=1)
            p_summary = doc.add_paragraph(
                "The demographic analysis reveals patterns that should inform timing, targeting, and messaging of future campaigns."
            )
            for run in p_summary.runs:
                run.font.name = 'Aptos'; run.font.size = Pt(9); run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

            doc_io = io.BytesIO(); doc.save(doc_io); doc_io.seek(0)
            st.success("Demographics executive summary ready.")
            st.download_button(
                "Download Demographics Executive Summary (.docx)",
                data=doc_io.getvalue(),
                file_name="Shopper_Demographics_Executive_Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Failed to generate Demographics document: {e}")

# ---------------- Main App ----------------
def main():
    st.title("PepsiCo Campaign Analysis Dashboard - 14 May 2025 to 14 August 2025")
    with st.sidebar:
        st.header("Navigation")
        menu = st.radio("Select Report", options=[
            "YOY Analysis",
            "Prior Periods",
            "Category Analysis",
            "Campaign Units Analysis",
            "Campaign Sales Amount Analysis",
            "Demographics"
        ])

    if menu == "YOY Analysis":
        yoy_analysis_page()
    elif menu == "Prior Periods":
        prior_periods_page()
    elif menu == "Category Analysis":
        category_analysis_page()
    elif menu == "Campaign Units Analysis":
        campaign_units_page()
    elif menu == "Campaign Sales Amount Analysis":
        campaign_sales_amount_page()
    elif menu == "Demographics":
        demographics_page()

if __name__ == "__main__":
    main()